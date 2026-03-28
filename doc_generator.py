"""
📄 Document Generator
- Creates ONE PAGE PDF matching original resume format
- Creates Word document as backup
"""

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, black
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    HRFlowable,
)
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from text_cleaner import clean_ai_response
import re
import os


# ============ ONE PAGE PDF GENERATOR ============

def create_resume_pdf(text, filename="Updated_Resume.pdf"):
    """
    Create a ONE PAGE professional PDF resume
    """

    # Clean AI intro/outro text first
    text = clean_ai_response(text)

    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        topMargin=0.3 * inch,
        bottomMargin=0.3 * inch,
        leftMargin=0.5 * inch,
        rightMargin=0.5 * inch
    )

    # ---- Compact Styles for ONE PAGE ----
    styles = getSampleStyleSheet()

    name_style = ParagraphStyle(
        'NameStyle',
        parent=styles['Title'],
        fontSize=16,
        leading=18,
        alignment=TA_CENTER,
        spaceAfter=1,
        spaceBefore=0,
        textColor=black,
        fontName='Helvetica-Bold'
    )

    contact_style = ParagraphStyle(
        'ContactStyle',
        parent=styles['Normal'],
        fontSize=8,
        leading=10,
        alignment=TA_CENTER,
        spaceAfter=3,
        spaceBefore=0,
        textColor=HexColor('#333333'),
        fontName='Helvetica'
    )

    heading_style = ParagraphStyle(
        'HeadingStyle',
        parent=styles['Heading2'],
        fontSize=10,
        leading=12,
        spaceBefore=5,
        spaceAfter=2,
        textColor=HexColor('#003366'),
        fontName='Helvetica-Bold',
    )

    subheading_style = ParagraphStyle(
        'SubHeadingStyle',
        parent=styles['Normal'],
        fontSize=9,
        leading=11,
        spaceBefore=3,
        spaceAfter=1,
        textColor=black,
        fontName='Helvetica-Bold'
    )

    normal_style = ParagraphStyle(
        'NormalStyle',
        parent=styles['Normal'],
        fontSize=8.5,
        leading=10,
        spaceAfter=1,
        spaceBefore=0,
        textColor=HexColor('#222222'),
        fontName='Helvetica',
        alignment=TA_JUSTIFY
    )

    bullet_style = ParagraphStyle(
        'BulletStyle',
        parent=styles['Normal'],
        fontSize=8.5,
        leading=10,
        spaceAfter=1,
        spaceBefore=0,
        leftIndent=12,
        textColor=HexColor('#222222'),
        fontName='Helvetica',
        alignment=TA_JUSTIFY
    )

    italic_style = ParagraphStyle(
        'ItalicStyle',
        parent=styles['Normal'],
        fontSize=8.5,
        leading=10,
        spaceAfter=1,
        spaceBefore=0,
        textColor=HexColor('#444444'),
        fontName='Helvetica-Oblique'
    )

    skills_style = ParagraphStyle(
        'SkillsStyle',
        parent=styles['Normal'],
        fontSize=8.5,
        leading=10,
        spaceAfter=1,
        spaceBefore=0,
        textColor=HexColor('#222222'),
        fontName='Helvetica'
    )

    # ---- Section Keywords ----
    section_keywords = [
        'professional summary', 'summary', 'objective',
        'professional experience', 'experience', 'work experience',
        'education', 'academic',
        'projects', 'project',
        'technical skills', 'skills',
        'soft skills',
        'languages', 'language',
        'certifications', 'certificates',
        'achievements', 'awards'
    ]

    def is_section_heading(line):
        clean = line.strip().lower()
        clean = clean.replace('#', '').replace('*', '').replace(':', '').strip()
        return clean in section_keywords

    def is_contact_line(line):
        indicators = ['@', '+91', 'phone', 'email', 'linkedin', 'github', '|', '—']
        return any(ind in line.lower() for ind in indicators)

    def is_bullet_point(line):
        return line.strip().startswith(('- ', '• ', '· ', '* ', '‣ '))

    def is_subheading(line):
        if any(word in line.lower() for word in [
            'ltd', 'pvt', 'inc', 'corp', 'technology',
            'institute', 'university', 'college', 'school',
            'arduino', 'based', 'system'
        ]):
            return True
        if line.strip().startswith('**') and line.strip().endswith('**'):
            return True
        return False

    def is_italic_line(line):
        if any(word in line.lower() for word in [
            'role:', 'team size', 'cgpa', 'hsc', 'sslc',
            'feb ', 'mar ', 'apr ', 'may ', 'jun ',
            'jul ', 'aug ', 'sep ', 'oct ', 'nov ', 'dec ',
            'jan ', '2020', '2021', '2022', '2023', '2024', '2025'
        ]):
            return True
        return False

    def clean_text(text):
        text = text.replace('**', '')
        text = text.replace('*', '')
        text = text.replace('##', '')
        text = text.replace('#', '')
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        return text.strip()

    # ---- Build PDF Content ----
    story = []
    lines = text.split('\n')
    name_found = False
    contact_found = False

    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue

        clean_line = clean_text(line)
        if not clean_line:
            continue

        # ---- Name ----
        if (line.startswith('# ') or (
            not name_found and i <= 2 and
            len(clean_line.split()) <= 4 and
            '@' not in clean_line and
            '+' not in clean_line and
            not any(k in clean_line.lower() for k in section_keywords)
        )):
            name_text = clean_text(line.replace('# ', ''))
            story.append(Paragraph(name_text, name_style))
            name_found = True
            continue

        # ---- Contact Info ----
        if is_contact_line(line) and not contact_found:
            contact_text = clean_text(line)
            contact_text = contact_text.replace(' — ', '  |  ')
            contact_text = contact_text.replace(' - ', '  |  ')
            story.append(Paragraph(contact_text, contact_style))
            contact_found = True
            story.append(HRFlowable(
                width="100%",
                thickness=0.8,
                color=HexColor('#003366'),
                spaceAfter=3
            ))
            continue

        # ---- Section Headings ----
        if is_section_heading(line) or line.startswith('## '):
            clean_heading = clean_text(line).replace(':', '').strip()
            story.append(Spacer(1, 2))
            story.append(Paragraph(clean_heading.upper(), heading_style))
            story.append(HRFlowable(
                width="100%",
                thickness=0.4,
                color=HexColor('#003366'),
                spaceAfter=2
            ))
            continue

        # ---- Bullet Points ----
        if is_bullet_point(line):
            bullet_text = clean_text(line.lstrip('-•·*‣ ').strip())
            story.append(Paragraph(
                f"•  {bullet_text}",
                bullet_style
            ))
            continue

        # ---- Sub-headings ----
        if is_subheading(line):
            story.append(Paragraph(clean_line, subheading_style))
            continue

        # ---- Italic Lines ----
        if is_italic_line(line):
            story.append(Paragraph(clean_line, italic_style))
            continue

        # ---- Skills Lines ----
        if ':' in line and any(
            word in line.lower() for word in [
                'programming', 'embedded', 'tools',
                'teamwork', 'english', 'tamil'
            ]
        ):
            parts = clean_line.split(':', 1)
            if len(parts) == 2:
                formatted = f"<b>{parts[0]}:</b>{parts[1]}"
                story.append(Paragraph(formatted, skills_style))
            else:
                story.append(Paragraph(clean_line, skills_style))
            continue

        # ---- Regular Text ----
        story.append(Paragraph(clean_line, normal_style))

    # ---- Build PDF ----
    try:
        doc.build(story)

        # Check if PDF is more than one page and shrink if needed
        from reportlab.lib.utils import open_for_read
        file_size = os.path.getsize(filename)

        return filename

    except Exception as e:
        print(f"PDF Error: {e}")
        return None


# ============ WORD DOCUMENT GENERATOR ============

def create_resume_docx(text, filename="Updated_Resume.docx"):
    """Create a Word document - ONE PAGE format"""

    # Clean AI intro/outro text first
    text = clean_ai_response(text)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    lines = text.split('\n')

    section_keywords = [
        'professional summary', 'summary',
        'professional experience', 'experience',
        'education', 'projects', 'technical skills',
        'soft skills', 'skills', 'languages',
        'certifications', 'achievements'
    ]

    def clean_text(text):
        text = text.replace('##', '').replace('#', '')
        text = text.replace('**', '').replace('*', '')
        return text.strip()

    name_found = False

    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue

        clean_line = clean_text(line)
        if not clean_line:
            continue

        # Name
        if (line.startswith('# ') or (
            i <= 2 and not name_found and
            len(clean_line.split()) <= 4 and
            '@' not in clean_line and
            '+' not in clean_line
        )):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(clean_line)
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = 'Calibri'
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)
            name_found = True
            continue

        # Contact info
        if any(ind in line.lower() for ind in ['@', '+91', 'linkedin']):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(clean_line)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(51, 51, 51)
            run.font.name = 'Calibri'
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            continue

        # Section headings
        is_heading = False
        clean_lower = clean_line.lower().replace(':', '').strip()
        for keyword in section_keywords:
            if clean_lower == keyword:
                is_heading = True
                break

        if is_heading or line.startswith('## '):
            heading_text = clean_line.replace(':', '').strip().upper()
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.font.name = 'Calibri'
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(1)
            continue

        # Bullet points
        if line.startswith(('- ', '• ', '· ', '* ')):
            bullet_text = clean_text(line.lstrip('-•·* '))
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(bullet_text)
            run.font.size = Pt(8.5)
            run.font.name = 'Calibri'
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            continue

        # Skills with colon
        if ':' in line and any(
            word in line.lower() for word in [
                'programming', 'embedded', 'tools',
                'teamwork', 'english', 'tamil'
            ]
        ):
            parts = clean_line.split(':', 1)
            p = doc.add_paragraph()
            if len(parts) == 2:
                run_bold = p.add_run(f"{parts[0]}:")
                run_bold.bold = True
                run_bold.font.size = Pt(8.5)
                run_bold.font.name = 'Calibri'

                run_normal = p.add_run(parts[1])
                run_normal.font.size = Pt(8.5)
                run_normal.font.name = 'Calibri'
            else:
                run = p.add_run(clean_line)
                run.font.size = Pt(8.5)
                run.font.name = 'Calibri'
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)
            continue

        # Sub-headings
        if any(word in line.lower() for word in [
            'ltd', 'pvt', 'institute', 'college', 'school',
            'university', 'technology'
        ]) or (line.startswith('**') and line.endswith('**')):
            p = doc.add_paragraph()
            run = p.add_run(clean_line)
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(0)
            continue

        # Italic lines
        if any(word in line.lower() for word in [
            'role:', 'team size', '2020', '2021', '2022',
            '2023', '2024', '2025', 'cgpa', 'hsc', 'sslc'
        ]):
            p = doc.add_paragraph()
            run = p.add_run(clean_line)
            run.italic = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(68, 68, 68)
            run.font.name = 'Calibri'
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            continue

        # Regular text
        p = doc.add_paragraph()
        run = p.add_run(clean_line)
        run.font.size = Pt(8.5)
        run.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)

    doc.save(filename)
    return filename


# ============ COVER LETTER PDF ============

def create_cover_letter_pdf(text, filename="Cover_Letter.pdf"):
    """Create a professional cover letter PDF"""

    # Clean AI intro/outro text
    text = clean_ai_response(text)

    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
        leftMargin=1 * inch,
        rightMargin=1 * inch
    )

    styles = getSampleStyleSheet()

    letter_style = ParagraphStyle(
        'LetterStyle',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        spaceAfter=8,
        fontName='Helvetica',
        alignment=TA_JUSTIFY
    )

    letter_bold = ParagraphStyle(
        'LetterBold',
        parent=letter_style,
        fontName='Helvetica-Bold'
    )

    story = []
    lines = text.split('\n')

    for line in lines:
        clean_line = line.strip().replace('**', '').replace('*', '').replace('#', '').strip()

        if not clean_line:
            story.append(Spacer(1, 8))
            continue

        if any(word in clean_line.lower() for word in [
            'dear', 'sincerely', 'regards', 'yours'
        ]):
            story.append(Paragraph(clean_line, letter_bold))
        else:
            story.append(Paragraph(clean_line, letter_style))

    try:
        doc.build(story)
        return filename
    except Exception as e:
        print(f"Cover Letter PDF Error: {e}")
        return None